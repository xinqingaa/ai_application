from __future__ import annotations

import json
import importlib.util
from pathlib import Path

import pytest
from docx import Document
from PIL import Image, ImageDraw
from reportlab.pdfgen import canvas

from rag_core import (
    ElementKind,
    EvidenceEligibility,
    IngestionError,
    IngestionErrorCode,
    IngestionStage,
    LoaderConfig,
    SourceRole,
    load_document,
)

REPO_ROOT = Path(__file__).resolve().parents[4]
FIXTURE_DIR = (
    REPO_ROOT / "source" / "apps" / "review_assistant" / "fixtures" / "rag" / "ingestion"
)


def _load(path: Path, **overrides):
    arguments = {
        "document_id": "KR-ORDER-STATE",
        "document_version": "1.0.0",
        "source_role": SourceRole.REFERENCE_KNOWLEDGE,
        "evidence_eligibility": EvidenceEligibility.CURRENT_EVIDENCE,
        "metadata": {"domain": "after_sale"},
    }
    arguments.update(overrides)
    return load_document(path, **arguments)


def test_txt_preserves_line_ranges_and_stable_identity(tmp_path: Path) -> None:
    path = tmp_path / "rules.txt"
    path.write_text("Current rules\npaid and completed\n\nAPI\nsource_channel required\n", encoding="utf-8")

    first = _load(path)
    second = _load(path)

    assert first.document.content_hash == second.document.content_hash
    assert [element.element_id for element in first.document.elements] == [
        element.element_id for element in second.document.elements
    ]
    assert first.document.metadata["domain"] == "after_sale"
    assert first.document.elements[0].locator.line_start == 1
    assert first.document.elements[0].locator.line_end == 2
    assert first.document.elements[1].locator.line_start == 4
    assert first.report.source_locator_kinds == ("text_lines",)


def test_markdown_preserves_heading_paths_and_list_items(tmp_path: Path) -> None:
    path = tmp_path / "rules.md"
    path.write_text(
        "# After-sale rules\n\n## States\n\n- paid\n- completed\n",
        encoding="utf-8",
    )

    result = _load(path)

    assert [element.kind for element in result.document.elements] == [
        ElementKind.TITLE,
        ElementKind.TITLE,
        ElementKind.LIST_ITEM,
        ElementKind.LIST_ITEM,
    ]
    assert result.document.elements[-1].locator.heading_path == (
        "After-sale rules",
        "States",
    )
    assert result.document.elements[-1].locator.line_start == 6


def test_docx_preserves_paragraph_and_table_locations(tmp_path: Path) -> None:
    path = tmp_path / "rules.docx"
    document = Document()
    document.add_heading("After-sale rules", level=1)
    document.add_paragraph("Only paid and completed orders are eligible.")
    table = document.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "field"
    table.cell(0, 1).text = "requirement"
    table.cell(1, 0).text = "source_channel"
    table.cell(1, 1).text = "required"
    document.save(str(path))

    result = _load(path)

    assert [element.kind for element in result.document.elements] == [
        ElementKind.TITLE,
        ElementKind.PARAGRAPH,
        ElementKind.TABLE,
    ]
    assert result.document.elements[1].locator.paragraph_index == 2
    assert result.document.elements[2].locator.table_index == 1
    assert result.document.elements[2].locator.heading_path == ("After-sale rules",)


def test_text_pdf_succeeds_with_page_location_and_warning(tmp_path: Path) -> None:
    path = tmp_path / "rules.pdf"
    pdf = canvas.Canvas(str(path))
    pdf.drawString(72, 760, "Only paid and completed orders are eligible.")
    pdf.showPage()
    pdf.drawString(72, 760, "source_channel is required by after-sale API v2.")
    pdf.save()

    result = _load(path)

    assert [element.locator.page_number for element in result.document.elements] == [1, 2]
    assert "paid and completed" in result.document.elements[0].text
    assert "source_channel" in result.document.elements[1].text
    assert {warning.code for warning in result.report.warnings} == {
        "pdf_reading_order_not_guaranteed"
    }


def test_pdf_without_text_layer_fails_explicitly(tmp_path: Path) -> None:
    image_path = tmp_path / "scan.png"
    image = Image.new("RGB", (600, 120), color="white")
    ImageDraw.Draw(image).text((20, 40), "SCANNED RULES", fill="black")
    image.save(image_path)
    path = tmp_path / "scan.pdf"
    pdf = canvas.Canvas(str(path))
    pdf.drawImage(str(image_path), 72, 680, width=300, height=60)
    pdf.showPage()
    pdf.save()

    with pytest.raises(IngestionError) as captured:
        _load(path)

    assert captured.value.stage is IngestionStage.EMPTY_CONTENT
    assert captured.value.code is IngestionErrorCode.PDF_TEXT_LAYER_MISSING


def test_pdf_with_partial_text_layer_reports_empty_page(tmp_path: Path) -> None:
    path = tmp_path / "partial.pdf"
    pdf = canvas.Canvas(str(path))
    pdf.drawString(72, 760, "text page")
    pdf.showPage()
    pdf.rect(72, 700, 300, 80, fill=1)
    pdf.showPage()
    pdf.save()

    result = _load(path)

    assert len(result.document.elements) == 1
    warnings = {warning.code: warning for warning in result.report.warnings}
    assert warnings["pdf_page_without_text"].locator is not None
    assert warnings["pdf_page_without_text"].locator.page_number == 2
    assert "pdf_reading_order_not_guaranteed" in warnings


def test_two_column_pdf_exposes_content_stream_reading_order() -> None:
    result = _load(FIXTURE_DIR / "reading_order_columns.pdf")

    assert result.document.text.index("右栏先写入") < result.document.text.index(
        "左栏应先读"
    )
    assert {
        warning.code for warning in result.report.warnings
    } == {"pdf_reading_order_not_guaranteed"}


def test_invalid_utf8_and_format_mismatch_have_different_stages(tmp_path: Path) -> None:
    invalid_text = tmp_path / "invalid.txt"
    invalid_text.write_bytes("售后".encode("gbk"))
    fake_pdf = tmp_path / "fake.pdf"
    fake_pdf.write_text("not a pdf", encoding="utf-8")

    with pytest.raises(IngestionError) as text_error:
        _load(invalid_text)
    with pytest.raises(IngestionError) as format_error:
        _load(fake_pdf)

    assert text_error.value.stage is IngestionStage.PARSE
    assert text_error.value.code is IngestionErrorCode.TEXT_DECODE_FAILED
    assert format_error.value.stage is IngestionStage.FORMAT_DETECTION
    assert format_error.value.code is IngestionErrorCode.FORMAT_MISMATCH


def test_empty_markdown_and_historical_evidence_boundary(tmp_path: Path) -> None:
    empty = tmp_path / "empty.md"
    empty.write_text("<!-- no knowledge content -->\n", encoding="utf-8")

    with pytest.raises(IngestionError) as empty_error:
        _load(empty)
    assert empty_error.value.code is IngestionErrorCode.EMPTY_DOCUMENT

    valid = tmp_path / "history.txt"
    valid.write_text("old review", encoding="utf-8")
    with pytest.raises(ValueError, match="Historical Material"):
        _load(
            valid,
            source_role=SourceRole.HISTORICAL_MATERIAL,
            evidence_eligibility=EvidenceEligibility.CURRENT_EVIDENCE,
        )

    history = _load(
        valid,
        source_role=SourceRole.HISTORICAL_MATERIAL,
        evidence_eligibility=EvidenceEligibility.HISTORICAL_CONTEXT,
    )
    assert history.document.source_role is SourceRole.HISTORICAL_MATERIAL
    assert history.document.evidence_eligibility is EvidenceEligibility.HISTORICAL_CONTEXT


def test_cleaning_actions_are_observable_and_preserve_semantics() -> None:
    result = load_document(
        FIXTURE_DIR / "cleaning_probe.md",
        document_id="CLEANING-PROBE",
        document_version="1.0.0",
        source_role=SourceRole.REFERENCE_KNOWLEDGE,
        evidence_eligibility=EvidenceEligibility.INELIGIBLE,
    )

    assert set(result.report.cleaning_actions) == {
        "replace_non_breaking_space",
        "normalize_unicode_nfc",
        "collapse_blank_lines",
        "trim_outer_whitespace",
    }
    assert "客户端 保留 Café 规则。" in result.document.text
    assert "line one\n\nline two" in result.document.text


def test_file_size_is_rejected_before_reading_bytes(
    tmp_path: Path,
    monkeypatch: pytest.MonkeyPatch,
) -> None:
    path = tmp_path / "large.txt"
    path.write_bytes(b"0123456789")

    def unexpected_read_bytes(_path: Path) -> bytes:
        raise AssertionError("超限文件不应进入 read_bytes")

    monkeypatch.setattr(Path, "read_bytes", unexpected_read_bytes)
    with pytest.raises(IngestionError) as captured:
        _load(path, config=LoaderConfig(max_file_bytes=5))

    assert captured.value.stage is IngestionStage.FORMAT_DETECTION
    assert captured.value.code is IngestionErrorCode.FILE_TOO_LARGE


def test_failure_manifest_freezes_stage_and_code() -> None:
    manifest = json.loads((FIXTURE_DIR / "manifest.json").read_text(encoding="utf-8"))

    for case in manifest["failure_cases"]:
        with pytest.raises(IngestionError) as captured:
            _load(FIXTURE_DIR / case["path"])
        assert captured.value.stage.value == case["expected_stage"]
        assert captured.value.code.value == case["expected_error"]


def test_failure_demo_treats_unexpected_success_as_mismatch(
    monkeypatch: pytest.MonkeyPatch,
) -> None:
    demo_path = REPO_ROOT / "source" / "demos" / "rag_ingestion_lab" / "inspect_ingestion.py"
    spec = importlib.util.spec_from_file_location("inspect_ingestion_test", demo_path)
    assert spec is not None and spec.loader is not None
    inspect_ingestion = importlib.util.module_from_spec(spec)
    spec.loader.exec_module(inspect_ingestion)

    monkeypatch.setattr(inspect_ingestion, "load_document", lambda *args, **kwargs: object())
    mismatches = inspect_ingestion._inspect_failures(
        [
            {
                "path": "empty.md",
                "expected_stage": "empty_content",
                "expected_error": "empty_document",
            }
        ],
        json_mode=True,
    )

    assert mismatches == 1


def test_controlled_formats_preserve_all_canonical_facts() -> None:
    manifest = json.loads((FIXTURE_DIR / "manifest.json").read_text(encoding="utf-8"))
    canonical = json.loads((FIXTURE_DIR / manifest["canonical_source"]).read_text(encoding="utf-8"))
    facts = [fact for section in canonical["sections"] for fact in section["facts"]]

    assert manifest["fixture_kind"] == "synthetic_controlled_format_comparison"
    assert manifest["comparison_mode"] == "mutually_exclusive_representations"

    for case in manifest["documents"]:
        result = load_document(
            FIXTURE_DIR / case["path"],
            document_id=case["document_id"],
            document_version=case["document_version"],
            source_role=SourceRole(case["source_role"]),
            evidence_eligibility=EvidenceEligibility(case["evidence_eligibility"]),
        )
        for fact in facts:
            assert fact in result.document.text, f"{case['path']} 丢失 canonical fact: {fact}"
