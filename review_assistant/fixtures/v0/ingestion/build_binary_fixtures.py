"""Rebuild the DOCX/PDF and stable failure fixtures used by the ingestion lab."""

from __future__ import annotations

import json
from pathlib import Path
from typing import Any

from docx import Document
from PIL import Image, ImageDraw, ImageFont
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.cidfonts import UnicodeCIDFont
from reportlab.pdfgen import canvas

FIXTURE_DIR = Path(__file__).resolve().parent
CANONICAL_PATH = FIXTURE_DIR / "canonical_content.json"
PDF_FONT = "STSong-Light"
CJK_FONT_CANDIDATES = (
    Path("/System/Library/Fonts/Supplemental/Arial Unicode.ttf"),
    Path("/System/Library/Fonts/STHeiti Light.ttc"),
    Path("/Library/Fonts/Arial Unicode.ttf"),
    Path("/usr/share/fonts/opentype/noto/NotoSansCJK-Regular.ttc"),
    Path("/usr/share/fonts/truetype/noto/NotoSansCJK-Regular.ttc"),
    Path("/usr/share/fonts/truetype/wqy/wqy-zenhei.ttc"),
)


def main() -> None:
    pdfmetrics.registerFont(UnicodeCIDFont(PDF_FONT))
    content = json.loads(CANONICAL_PATH.read_text(encoding="utf-8"))
    _build_docx(FIXTURE_DIR / "order_rules.docx", content)
    _build_text_pdf(FIXTURE_DIR / "order_rules.pdf", content)
    _build_image_only_pdf(FIXTURE_DIR / "image_only_scan.pdf")
    _build_reading_order_pdf(FIXTURE_DIR / "reading_order_columns.pdf")
    (FIXTURE_DIR / "damaged.docx").write_bytes(b"PK damaged OOXML fixture")
    (FIXTURE_DIR / "invalid_encoding.txt").write_bytes("售后入口".encode("gbk"))


def _build_docx(path: Path, content: dict[str, Any]) -> None:
    document = Document()
    document.add_heading(content["title"], level=1)
    current_rules, constraints = content["sections"]
    document.add_heading(current_rules["heading"], level=2)
    for fact in current_rules["facts"]:
        document.add_paragraph(fact)
    document.add_heading(constraints["heading"], level=2)
    table = document.add_table(rows=3, cols=2)
    table.cell(0, 0).text = "对象"
    table.cell(0, 1).text = "当前约束"
    table.cell(1, 0).text = "售后接口 v2"
    table.cell(1, 1).text = constraints["facts"][0]
    table.cell(2, 0).text = "Flutter 客户端"
    table.cell(2, 1).text = constraints["facts"][1]
    document.save(str(path))


def _build_text_pdf(path: Path, content: dict[str, Any]) -> None:
    pdf = canvas.Canvas(str(path))
    pdf.setTitle(content["title"])
    pdf.setFont(PDF_FONT, 16)
    pdf.drawString(72, 760, content["title"])
    pdf.setFont(PDF_FONT, 11)
    current_rules, constraints = content["sections"]
    pdf.drawString(72, 730, current_rules["heading"])
    for offset, fact in enumerate(current_rules["facts"]):
        pdf.drawString(72, 705 - offset * 20, fact)
    pdf.showPage()
    pdf.setFont(PDF_FONT, 13)
    pdf.drawString(72, 760, constraints["heading"])
    pdf.setFont(PDF_FONT, 11)
    for offset, fact in enumerate(constraints["facts"]):
        pdf.drawString(72, 730 - offset * 20, fact)
    pdf.save()


def _build_image_only_pdf(path: Path) -> None:
    image_path = FIXTURE_DIR / "_scan_source.png"
    image = Image.new("RGB", (880, 240), color="white")
    drawing = ImageDraw.Draw(image)
    drawing.rectangle((5, 5, 875, 235), outline="black", width=4)
    font = _load_cjk_image_font(size=36)
    drawing.text((40, 95), "扫描件：售后规则", fill="black", font=font)
    image.save(image_path)
    pdf = canvas.Canvas(str(path))
    pdf.setTitle("无文本层扫描 PDF fixture")
    pdf.drawImage(str(image_path), 72, 620, width=440, height=120)
    pdf.showPage()
    pdf.save()
    image_path.unlink()


def _build_reading_order_pdf(path: Path) -> None:
    pdf = canvas.Canvas(str(path))
    pdf.setTitle("双栏阅读顺序对照 fixture")
    pdf.setFont(PDF_FONT, 12)
    # Intentionally write the right column first in the content stream.
    pdf.drawString(320, 760, "右栏先写入")
    pdf.drawString(72, 760, "左栏应先读")
    pdf.setFont(PDF_FONT, 10)
    pdf.drawString(320, 735, "抽取结果会先出现这一列。")
    pdf.drawString(72, 735, "视觉上这一列应先阅读。")
    pdf.showPage()
    pdf.save()


def _load_cjk_image_font(*, size: int):
    for candidate in CJK_FONT_CANDIDATES:
        if not candidate.exists():
            continue
        try:
            return ImageFont.truetype(str(candidate), size=size)
        except OSError:
            continue
    raise FileNotFoundError(
        "重建扫描 PDF 需要本地 CJK 字体；常见路径均未找到。"
        "可安装 Noto Sans CJK 后重试。"
    )


if __name__ == "__main__":
    main()
